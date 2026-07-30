from docx import Document
from db.db_operations import insert_data, delete_data
from backend.config import Config
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from utils.logging_util import LoggerSingleton
import logging
import subprocess
import tempfile
import os
import asyncio
from typing import List
from edu.file_text_extract_timeout import await_file_text_extract

# from db.db_redis import job_progress_manager, job_manager
from db.db_job_managers import AsyncJobManager, AsyncJobProgress


# 로거 설정
logger = LoggerSingleton.get_logger(logger_name="edu.doc", level=logging.INFO)

embedding_model = OpenAIEmbeddings(openai_api_key=Config.OPENAI_API_KEY)
FILE_EMBEDDING_BATCH_SIZE = max(1, int(getattr(Config, "FILE_EMBEDDING_BATCH_SIZE", 5) or 5))


# ✅ DOC 병렬 처리 함수들 추가 (HWP 처리와 동일한 패턴)
async def process_single_chunk_async_doc(
    chunk: str,
    chunk_idx: int,
    content: str,
    subject: str,
    table_name: str,
    dbname: str,
    memo: str,
    job_id: str,
    job_manager: AsyncJobManager,
    job_progress_manager: AsyncJobProgress = None,
    chunk_progress: float = 0.0,
    personal_info_filter: str = "N",  # 개인정보 필터링 옵션 추가
):
    """개별 DOC 청크를 비동기로 처리 (임베딩 + 저장, 배치에서 진행률 관리)"""
    try:
        # ✅ 청크 처리 전 취소 상태 확인
        status = await job_manager.get_job_status(job_id)
        if status == "cancel":
            logger.info(f"DOC 개별 청크 처리 중 작업 취소됨: job_id={job_id}, chunk={chunk_idx}")
            return {"status": "cancelled", "chunk_idx": chunk_idx}
        personal_info_filter = personal_info_filter
        if personal_info_filter == "Y":
            logger.info(f"DOC 개별 청크 처리 중 개인정보 필터링 적용: job_id={job_id}, chunk={chunk_idx}")
            from utils.dlp_api import check_pii_content
            original_chunk = chunk  # 원본 청크 저장
            pii_result = check_pii_content(chunk)
            
            # 마스킹된 텍스트를 chunk 변수에 할당
            if pii_result["success"]:
                chunk = pii_result["masked_text"]
                masked_parts_text = pii_result.get("masked_parts_text", "")

            else:
                # PII 검사 실패 시 에러 로깅
                logger.error(f"PII 검사 실패: {pii_result.get('error', 'Unknown error')}")
                # 실패 시 원본 텍스트 유지
        chunk_num = str(chunk_idx)
        chunk_with_metadata = chunk  # 본문만 저장
        
        # ✅ 비동기 임베딩 생성
        embedding = await embedding_model.aembed_query(chunk_with_metadata)
        embedding_array = f"[{','.join(map(str, embedding))}]"

        # ✅ DB 저장 전 한 번 더 취소 상태 확인
        status = await job_manager.get_job_status(job_id)
        if status == "cancel":
            logger.info(f"DOC DB 저장 전 작업 취소됨: job_id={job_id}, chunk={chunk_idx}")
            return {"status": "cancelled", "chunk_idx": chunk_idx}

        # ✅ 비동기 DB 저장
        await insert_data(
            table=table_name,
            data={
                "content": content,
                "subject": subject,
                "chunk_num": chunk_num,
                "memo": memo,
                "content_type": "file",
                "text_data": chunk_with_metadata,
                "embedding": embedding_array,
            },
            dbname=dbname,
        )

        # ✅ 개별 청크에서는 진행률을 업데이트하지 않음 (배치에서 관리)
        logger.debug(f"[DOC 청크 처리 완료] job_id={job_id}, chunk={chunk_idx}")

        return {"status": "success", "chunk_idx": chunk_idx}

    except Exception as e:
        logger.error(f"DOC 청크 처리 오류 ({content}, chunk {chunk_idx}): {e}")
        return {"status": "error", "chunk_idx": chunk_idx, "error": str(e)}


async def process_chunks_parallel_doc(
    chunks: List[str],
    content: str,
    subject: str,
    table_name: str,
    dbname: str,
    job_id: str,
    job_manager: AsyncJobManager,
    memo: str,
    job_progress_manager: AsyncJobProgress = None,
    chunk_progress: float = 0.0,
    batch_size: int = 5,
    personal_info_filter: str = "N",  # 개인정보 필터링 옵션 추가
):
    """DOC 청크들을 배치 단위로 병렬 처리 (TXT와 동일한 진행률 관리)"""
    
    total_chunks = len(chunks)
    completed_chunks = 0
    
    logger.info(f"[DOC 배치 병렬 처리 시작] 파일: {content}, 총 청크: {total_chunks}, 배치 크기: {batch_size}")
    
    # 배치별로 청크 처리
    for i in range(0, len(chunks), batch_size):
        # 취소 확인
        status = await job_manager.get_job_status(job_id)
        if status == "cancel":
            logger.info(f"DOC 청크 처리 중 작업 취소됨: job_id={job_id}")
            return

        batch_chunks = chunks[i:i + batch_size]
        batch_tasks = []

        # 배치 내 병렬 처리
        for chunk_idx, chunk in enumerate(batch_chunks, start=i + 1):
            task = process_single_chunk_async_doc(
                chunk=chunk,
                chunk_idx=chunk_idx,
                content=content,
                subject=subject,
                table_name=table_name,
                dbname=dbname,
                memo=memo,
                job_id=job_id,
                job_manager=job_manager,
                job_progress_manager=job_progress_manager,
                chunk_progress=chunk_progress,
                personal_info_filter=personal_info_filter, # 개인정보 필터링 옵션 전달
            )
            batch_tasks.append(task)

        # ✅ 배치 실행 및 진행률 업데이트 (TXT와 동일한 로직)
        try:
            results = await asyncio.gather(*batch_tasks, return_exceptions=True)
            
            # 배치에서 성공한 청크 수만큼 진행률 업데이트
            successful_chunks = 0
            for result in results:
                if not isinstance(result, Exception) and isinstance(result, dict) and result.get("status") == "success":
                    successful_chunks += 1
            
            if successful_chunks > 0:
                completed_chunks += successful_chunks
                
                # ✅ 배치 완료 시 한 번에 진행률 업데이트
                current_progress = await job_progress_manager.get_job_progress(job_id)
                progress_increment = chunk_progress * successful_chunks
                new_progress = round(min(current_progress + progress_increment, 99.99), 2)
                
                logger.info(f"[DOC 진행률 업데이트] 파일: {content}, "
                           f"성공 청크: {successful_chunks}/{len(batch_chunks)}, "
                           f"청크별 진행률: {chunk_progress:.6f}%, "
                           f"증가량: {progress_increment:.6f}%, "
                           f"현재 진행률: {current_progress}% → {new_progress}%, "
                           f"완료된 총 청크: {completed_chunks}/{total_chunks}")
                
                await job_progress_manager.set_job_progress(job_id, new_progress)
                

            
            logger.debug(f"DOC 배치 처리 완료: {content}, 배치 {i//batch_size + 1}, "
                        f"성공 청크: {successful_chunks}/{len(batch_chunks)}")
            
        except Exception as e:
            logger.error(f"DOC 배치 처리 중 오류: {e}")

    logger.info(f"[DOC 파일 청크 처리 완료] {content}: {completed_chunks}/{total_chunks} 청크 성공")


async def process_doc(
    content: str,
    file_path: str,
    table_name: str,
    dbname: str,
    job_id: str,
    each_progress: float,
    job_manager: AsyncJobManager,
    job_progress_manager: AsyncJobProgress,
    subject: str = "",
    memo: str = "",
    personal_info_filter: str = "N",  # 개인정보 필터링 옵션 추가
):
    """
    DOC 파일을 병렬 처리하여 텍스트를 추출하고 벡터화한 후 데이터베이스에 저장합니다.
    """
    try:
        _, ext = os.path.splitext(file_path)

        if ext.lower() == ".doc":
            # ✅ 비동기로 antiword 실행
            try:
                result = await await_file_text_extract(
                    asyncio.to_thread(
                        subprocess.run,
                        ["antiword", file_path],
                        capture_output=True,
                        text=True,
                    ),
                    path=file_path,
                    stage="doc",
                    logger=logger,
                )
                all_text = result.stdout
            except subprocess.CalledProcessError as e:
                raise RuntimeError(f"Antiword 처리 중 오류 발생: {e}")
        else:
            # ✅ 비동기로 .docx 처리
            def extract_docx_content(file_path):
                doc = Document(file_path)
                all_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                # 표 처리
                for table in doc.tables:
                    table_text = table_to_markdown(table)
                    all_text += "\n" + table_text
                
                return all_text
            
            all_text = await await_file_text_extract(
                asyncio.to_thread(extract_docx_content, file_path),
                path=file_path,
                stage="docx",
                logger=logger,
            )

        # 텍스트 분할 및 임베딩 처리
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.BASIC_CHUNK_SIZE, chunk_overlap=Config.BASIC_CHUNK_OVERLAP
        )
        chunks = text_splitter.split_text(all_text)
        total_chunks = len(chunks)

        if total_chunks == 0:
            logger.error(f"[content: {content}] 분할된 청크가 없습니다. 파일 내용이 비어있거나 분할에 실패했습니다. job_id:{job_id}, table_name: {table_name}")
            total_chunks = 1
            chunks = ["[이 문서는 텍스트를 포함하지 않습니다.]"]

        # ✅ DOC 파일: 이 파일의 올바른 청크별 진행률 계산
        # 단일 파일 처리라고 가정하고 95%를 총 청크 수로 나눔 (FAISS 5% 제외)
        chunk_progress = round(95 / total_chunks, 6)
        
        logger.info(
            f"[DOC 병렬 처리 시작] content: {content}, total_chunk: {total_chunks}, "
            f"청크별 진행률: {chunk_progress}%, job_id: {job_id}"
        )

        # ✅ 청크들을 배치 단위로 병렬 처리 (TXT와 동일한 진행률 관리)
        await process_chunks_parallel_doc(
            chunks=chunks,
            content=content,
            subject=subject or os.path.basename(file_path or "") or content,
            table_name=table_name,
            dbname=dbname,
            job_id=job_id,
            job_manager=job_manager,
            memo=memo,
            job_progress_manager=job_progress_manager,
            chunk_progress=chunk_progress,
            batch_size=FILE_EMBEDDING_BATCH_SIZE,
            personal_info_filter=personal_info_filter, # 개인정보 필터링 옵션 전달
        )

        logger.info(
            f"[DOC 병렬 처리 완료] content: {content}, total_chunk: {total_chunks}, job_id: {job_id}"
        )
        return {
            "status": "success", 
            "message": f"{content} 처리 완료", 
            "chunks": total_chunks,
            "chunk_count": [total_chunks],
            "use_source": [content]
        }
    except Exception as e:
        logger.error(f"DOC 처리 중 오류 발생: {e}", exc_info=True)

        raise RuntimeError(f"DOC 처리 중 오류 발생: {e}")


def table_to_markdown(table):
    """
    Word 문서의 표를 마크다운 형식으로 변환합니다.
    """
    markdown_rows = []

    # 헤더 행 처리
    header_cells = []
    for cell in table.rows[0].cells:
        header_cells.append(cell.text.strip())
    markdown_rows.append("| " + " | ".join(header_cells) + " |")

    # 구분선 추가
    markdown_rows.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

    # 데이터 행 처리
    for row in table.rows[1:]:
        row_cells = []
        for cell in row.cells:
            row_cells.append(cell.text.strip())
        markdown_rows.append("| " + " | ".join(row_cells) + " |")

    return "\n".join(markdown_rows)


def extract_doc_plain_text_sync(file_path: str) -> str:
    """process_doc과 동일 규칙으로 본문만 추출(청크·임베딩·DB 없음)."""
    if not file_path or not os.path.isfile(file_path):
        return ""
    _, ext = os.path.splitext(file_path)
    if ext.lower() == ".doc":
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True,
            text=True,
        )
        return (result.stdout or "").strip()
    doc = Document(file_path)
    all_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    for table in doc.tables:
        all_text += "\n" + table_to_markdown(table)
    return (all_text or "").strip()


async def calculate_doc_chunks(file_path: str) -> int:
    """DOC 파일의 정확한 청크 수를 계산합니다."""
    try:
        import asyncio
        import subprocess
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from backend.config import Config
        from docx import Document
        
        _, ext = os.path.splitext(file_path)
        
        if ext.lower() == ".doc":
            # antiword로 텍스트 추출
            def extract_doc_text(file_path):
                result = subprocess.run(
                    ["antiword", file_path], 
                    capture_output=True, 
                    text=True
                )
                return result.stdout
            
            all_text = await asyncio.to_thread(extract_doc_text, file_path)
        else:
            # .docx 파일 처리
            def extract_docx_text(file_path):
                doc = Document(file_path)
                all_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                # 표 처리도 포함
                for table in doc.tables:
                    table_text = table_to_markdown(table)
                    all_text += "\n" + table_text
                
                return all_text
            
            all_text = await asyncio.to_thread(extract_docx_text, file_path)
        
        if not all_text.strip():
            logger.warning(f"DOC 파일에서 텍스트를 추출할 수 없습니다: {file_path}")
            return 1  # 기본값
        
        # ✅ 동일한 청크 분할 설정 사용
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.BASIC_CHUNK_SIZE, 
            chunk_overlap=Config.BASIC_CHUNK_OVERLAP
        )
        chunks = text_splitter.split_text(all_text)
        actual_chunks = len(chunks)
        
        logger.info(f"DOC 정확한 청크 수 계산: {file_path}, 실제 청크: {actual_chunks}개")
        return actual_chunks
        
    except Exception as e:
        logger.error(f"DOC 청크 수 계산 실패: {file_path}, 오류: {e}")
        return 5  # 기본값
