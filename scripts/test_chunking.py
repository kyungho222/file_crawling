"""
파일 일부만 청킹 테스트 스크립트

이 스크립트는 다음 기능을 제공합니다:
1. 특정 파일 또는 폴더의 파일들에 대해 청킹 테스트
2. PDF, HWP, TEXT, DOC 등 다양한 형식 지원
3. 전체 파일이 아닌 일부만 처리 (페이지 수 제한, 크기 제한 등)
4. 청크 결과를 파일로 저장 및 콘솔 출력

사용법:
    python scripts/test_chunking.py --file <파일경로> --max-chunks 10
    python scripts/test_chunking.py --dir downloads --file-type pdf --max-pages 5
"""

import asyncio
import os
import sys
from pathlib import Path
import argparse
from typing import List, Dict, Any
import json

# 프로젝트 루트를 sys.path에 추가
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from langchain_text_splitters import RecursiveCharacterTextSplitter
from config.settings import Config
from langchain_openai import OpenAIEmbeddings

# 파일 타입별 텍스트 추출 함수
async def extract_text_from_pdf(file_path: str, max_pages: int = None) -> str:
    """PDF에서 텍스트 추출 (페이지 제한 가능)"""
    try:
        import pdfplumber
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            pages_to_process = min(max_pages or total_pages, total_pages)
            
            print(f"  📄 PDF 총 페이지: {total_pages}, 처리할 페이지: {pages_to_process}")
            
            for i in range(pages_to_process):
                page = pdf.pages[i]
                text = page.extract_text()
                if text:
                    text_parts.append(f"[페이지 {i+1}]\n{text}")
        
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"  ❌ PDF 추출 오류: {e}")
        return ""


async def extract_text_from_hwp(file_path: str, max_pages: int = None) -> str:
    """HWP에서 텍스트 추출 (페이지 제한 가능)"""
    try:
        import subprocess
        import shutil
        import importlib.util
        import sys
        
        def detect_hwp_version(file_path):
            """파일의 HWP 버전을 감지하는 함수"""
            with open(file_path, "rb") as f:
                header = f.read(8)
            
            if header.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
                return "HWP5"
            elif header[:4] == b"HWP3":
                return "HWP3"
            elif header[:4] == b"HWP2":
                return "HWP2"
            else:
                return None
        
        def build_hwp_txt_command(version: str, hwp_file_path: str):
            """환경에 따라 HWP 텍스트 변환 커맨드를 구성"""
            exe_map = {"HWP5": "hwp5txt", "HWP3": "hwp3txt", "HWP2": "hwp2txt"}
            module_map = {"HWP5": "hwp5.hwp5txt", "HWP3": "hwp5.hwp3txt", "HWP2": "hwp5.hwp2txt"}
            
            exe = exe_map.get(version)
            if not exe:
                raise ValueError("지원하지 않는 HWP 파일 형식입니다.")
            
            if shutil.which(exe):
                return [exe, hwp_file_path]
            
            mod = module_map.get(version)
            if mod and importlib.util.find_spec(mod) is not None:
                return [sys.executable, "-m", mod, hwp_file_path]
            
            if version == "HWP5" and importlib.util.find_spec("hwp5txt") is not None:
                return [sys.executable, "-m", "hwp5txt", hwp_file_path]
            
            raise FileNotFoundError(f"{exe} 실행 파일을 찾을 수 없습니다. pip install hwp5 를 실행하세요.")
        
        def hwp_to_text_sync(hwp_file_path):
            """HWP 파일을 텍스트로 변환하는 함수"""
            if not os.path.isfile(hwp_file_path):
                raise FileNotFoundError(f"파일을 찾을 수 없습니다: {hwp_file_path}")
            
            version = detect_hwp_version(hwp_file_path)
            if version not in ("HWP5", "HWP3", "HWP2"):
                raise ValueError("지원하지 않는 HWP 파일 형식입니다.")
            
            cmd = build_hwp_txt_command(version, hwp_file_path)
            result = subprocess.run(
                cmd,
                shell=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                check=True,
            )
            return result.stdout
        
        # hwp_to_text를 비동기로 실행
        full_text = await asyncio.to_thread(hwp_to_text_sync, file_path)
        
        # max_pages 제한이 있는 경우 텍스트 일부만 반환
        if max_pages and full_text:
            lines = full_text.split('\n')
            lines_per_page = 50
            max_lines = max_pages * lines_per_page
            limited_text = '\n'.join(lines[:max_lines])
            
            print(f"  📄 HWP 전체 줄 수: {len(lines)}, 제한 줄 수: {max_lines} (약 {max_pages}페이지)")
            return limited_text
        
        return full_text
    except Exception as e:
        print(f"  ❌ HWP 추출 오류: {e}")
        return ""


async def extract_text_from_txt(file_path: str, max_chars: int = None) -> str:
    """텍스트 파일에서 텍스트 추출 (문자 수 제한 가능)"""
    try:
        encodings = ['utf-8', 'cp949', 'euc-kr', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read(max_chars) if max_chars else f.read()
                    
                if max_chars:
                    print(f"  📝 텍스트 파일: 최대 {max_chars}자까지 읽음")
                
                return text
            except UnicodeDecodeError:
                continue
        
        print(f"  ❌ 텍스트 파일 인코딩을 찾을 수 없습니다")
        return ""
    except Exception as e:
        print(f"  ❌ 텍스트 추출 오류: {e}")
        return ""


async def extract_text_from_doc(file_path: str, max_paragraphs: int = None) -> str:
    """DOC/DOCX에서 텍스트 추출 (단락 수 제한 가능)"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        paragraphs = []
        
        total_paragraphs = len(doc.paragraphs)
        paragraphs_to_process = min(max_paragraphs or total_paragraphs, total_paragraphs)
        
        print(f"  📄 DOC 총 단락: {total_paragraphs}, 처리할 단락: {paragraphs_to_process}")
        
        for i in range(paragraphs_to_process):
            para = doc.paragraphs[i]
            if para.text.strip():
                paragraphs.append(para.text)
        
        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"  ❌ DOC 추출 오류: {e}")
        return ""


async def perform_chunking(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    max_chunks: int = None
) -> List[str]:
    """텍스트를 청크로 분할"""
    
    if not text or not text.strip():
        print(f"  ⚠️  추출된 텍스트가 비어있습니다")
        return []
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
    )
    
    chunks = text_splitter.split_text(text)
    
    if max_chunks:
        chunks = chunks[:max_chunks]
        print(f"  ✂️  청크 제한: {max_chunks}개까지만 처리")
    
    return chunks


async def test_file_chunking(
    file_path: str,
    max_pages: int = None,
    max_chars: int = None,
    max_chunks: int = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    output_dir: str = None
) -> Dict[str, Any]:
    """단일 파일에 대한 청킹 테스트"""
    
    file_path = Path(file_path)
    
    if not file_path.exists():
        print(f"❌ 파일을 찾을 수 없습니다: {file_path}")
        return None
    
    print(f"\n{'='*60}")
    print(f"📁 파일: {file_path.name}")
    print(f"{'='*60}")
    
    # 파일 확장자에 따라 텍스트 추출
    ext = file_path.suffix.lower()
    
    if ext == '.pdf':
        text = await extract_text_from_pdf(str(file_path), max_pages)
    elif ext in ['.hwp', '.hwpx']:
        text = await extract_text_from_hwp(str(file_path), max_pages)
    elif ext in ['.txt', '.text']:
        text = await extract_text_from_txt(str(file_path), max_chars)
    elif ext in ['.doc', '.docx']:
        text = await extract_text_from_doc(str(file_path), max_pages)
    else:
        print(f"  ⚠️  지원하지 않는 파일 형식: {ext}")
        return None
    
    if not text:
        print(f"  ❌ 텍스트 추출 실패")
        return None
    
    # 텍스트 정보 출력
    print(f"  📊 추출된 텍스트 길이: {len(text):,}자")
    print(f"  📊 줄 수: {text.count(chr(10)):,}줄")
    
    # 청킹 수행
    print(f"\n  🔪 청킹 시작 (chunk_size={chunk_size}, overlap={chunk_overlap})...")
    chunks = await perform_chunking(text, chunk_size, chunk_overlap, max_chunks)
    
    print(f"  ✅ 생성된 청크 수: {len(chunks)}개")
    
    # 결과 정보
    result = {
        "file_name": file_path.name,
        "file_path": str(file_path),
        "file_type": ext,
        "text_length": len(text),
        "text_lines": text.count('\n'),
        "chunk_count": len(chunks),
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "chunks": chunks
    }
    
    # 청크 미리보기
    print(f"\n  {'─'*58}")
    print(f"  📋 청크 미리보기 (처음 3개):")
    print(f"  {'─'*58}")
    for i, chunk in enumerate(chunks[:3]):
        print(f"\n  [청크 {i+1}/{len(chunks)}] ({len(chunk)}자)")
        preview = chunk[:200].replace('\n', ' ')
        if len(chunk) > 200:
            preview += "..."
        print(f"  {preview}")
    
    # 결과 저장
    if output_dir:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        # JSON 결과 저장
        json_file = output_path / f"{file_path.stem}_chunks.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump({
                **result,
                "chunks": [{"index": i, "length": len(c), "text": c} for i, c in enumerate(chunks)]
            }, f, ensure_ascii=False, indent=2)
        
        # 텍스트 결과 저장
        txt_file = output_path / f"{file_path.stem}_chunks.txt"
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(f"파일: {file_path.name}\n")
            f.write(f"총 청크 수: {len(chunks)}\n")
            f.write(f"{'='*60}\n\n")
            for i, chunk in enumerate(chunks):
                f.write(f"[청크 {i+1}/{len(chunks)}] ({len(chunk)}자)\n")
                f.write(f"{'-'*60}\n")
                f.write(chunk)
                f.write(f"\n\n{'='*60}\n\n")
        
        print(f"\n  💾 결과 저장:")
        print(f"     - JSON: {json_file}")
        print(f"     - TXT:  {txt_file}")
    
    return result


async def main():
    parser = argparse.ArgumentParser(
        description="파일 일부만 청킹 테스트",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
예제:
  # 단일 PDF 파일의 처음 5페이지만 청킹
  python scripts/test_chunking.py --file downloads/sample.pdf --max-pages 5
  
  # 텍스트 파일의 처음 10000자만 청킹
  python scripts/test_chunking.py --file downloads/sample.txt --max-chars 10000
  
  # 폴더 내 모든 PDF 파일 청킹 (각각 3페이지, 최대 10청크)
  python scripts/test_chunking.py --dir downloads --file-type pdf --max-pages 3 --max-chunks 10
  
  # 결과를 output 폴더에 저장
  python scripts/test_chunking.py --file downloads/sample.pdf --output output
        """
    )
    
    parser.add_argument('--file', type=str, help='테스트할 파일 경로')
    parser.add_argument('--dir', type=str, default='downloads', help='테스트할 폴더 경로 (기본: downloads)')
    parser.add_argument('--file-type', type=str, choices=['pdf', 'hwp', 'txt', 'doc', 'all'], default='all',
                        help='처리할 파일 타입 (기본: all)')
    parser.add_argument('--max-pages', type=int, help='처리할 최대 페이지 수 (PDF, HWP, DOC)')
    parser.add_argument('--max-chars', type=int, help='처리할 최대 문자 수 (TXT)')
    parser.add_argument('--max-chunks', type=int, help='생성할 최대 청크 수')
    parser.add_argument('--chunk-size', type=int, default=1000, help='청크 크기 (기본: 1000)')
    parser.add_argument('--chunk-overlap', type=int, default=200, help='청크 겹침 크기 (기본: 200)')
    parser.add_argument('--output', type=str, help='결과 저장 폴더')
    
    args = parser.parse_args()
    
    print("\n" + "="*60)
    print("🧪 파일 청킹 테스트 시작")
    print("="*60)
    
    results = []
    
    # 단일 파일 처리
    if args.file:
        result = await test_file_chunking(
            file_path=args.file,
            max_pages=args.max_pages,
            max_chars=args.max_chars,
            max_chunks=args.max_chunks,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
            output_dir=args.output
        )
        if result:
            results.append(result)
    
    # 폴더 내 파일들 처리
    else:
        dir_path = Path(args.dir)
        
        if not dir_path.exists():
            print(f"❌ 폴더를 찾을 수 없습니다: {dir_path}")
            return
        
        # 파일 타입 필터
        extensions = {
            'pdf': ['.pdf'],
            'hwp': ['.hwp', '.hwpx'],
            'txt': ['.txt', '.text'],
            'doc': ['.doc', '.docx'],
            'all': ['.pdf', '.hwp', '.hwpx', '.txt', '.text', '.doc', '.docx']
        }
        
        target_extensions = extensions.get(args.file_type, extensions['all'])
        
        # 파일 목록 가져오기
        files = [f for f in dir_path.iterdir() if f.is_file() and f.suffix.lower() in target_extensions]
        
        if not files:
            print(f"⚠️  {dir_path}에서 처리할 파일을 찾을 수 없습니다 (타입: {args.file_type})")
            return
        
        print(f"\n📂 폴더: {dir_path}")
        print(f"📄 발견된 파일: {len(files)}개")
        
        for file_path in files:
            result = await test_file_chunking(
                file_path=str(file_path),
                max_pages=args.max_pages,
                max_chars=args.max_chars,
                max_chunks=args.max_chunks,
                chunk_size=args.chunk_size,
                chunk_overlap=args.chunk_overlap,
                output_dir=args.output
            )
            if result:
                results.append(result)
    
    # 전체 결과 요약
    print(f"\n{'='*60}")
    print(f"📊 전체 테스트 결과 요약")
    print(f"{'='*60}")
    print(f"  처리된 파일 수: {len(results)}개")
    
    if results:
        total_chunks = sum(r['chunk_count'] for r in results)
        total_text_length = sum(r['text_length'] for r in results)
        print(f"  총 생성 청크 수: {total_chunks}개")
        print(f"  총 텍스트 길이: {total_text_length:,}자")
        
        print(f"\n  파일별 청크 수:")
        for r in results:
            print(f"    - {r['file_name']}: {r['chunk_count']}개 청크")
    
    print(f"\n✅ 테스트 완료!\n")


if __name__ == "__main__":
    asyncio.run(main())
